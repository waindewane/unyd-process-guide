#!/usr/bin/env python3
"""Generate minimally formatted, process-specific DOCX prose exports."""

from __future__ import annotations

import json
import sys
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT = "Times New Roman"
BODY_SIZE = Pt(11)
TITLE_SIZE = Pt(12)
LINK_BLUE = "0563C1"


def set_run_font(run, size=BODY_SIZE, bold=None, italic=None, color=None):
    run.font.name = FONT
    run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    properties.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    properties.append(size)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_BLUE)
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_text(paragraph, text, *, bold=False, italic=False):
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold, italic=italic)
    return run


def add_labeled_paragraph(document, label, text):
    paragraph = document.add_paragraph()
    add_text(paragraph, f"{label}: ", bold=True)
    add_text(paragraph, text)
    return paragraph


def add_bullet(document, text, *, label=None, href=None, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = document.add_paragraph(style=style)
    if label:
        if href:
            add_hyperlink(paragraph, label, href)
            add_text(paragraph, ": ")
        else:
            add_text(paragraph, f"{label}: ", bold=True)
    add_text(paragraph, text)
    return paragraph


def add_source(document, label, href, detail=None):
    paragraph = document.add_paragraph(style="List Bullet")
    add_hyperlink(paragraph, label, href)
    if detail:
        add_text(paragraph, f": {detail}")


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY_SIZE
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.08

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = FONT
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.bold = True
        style.font.size = TITLE_SIZE if style_name == "Title" else BODY_SIZE
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.space_before = Pt(0 if style_name == "Title" else 9)
        paragraph_properties = style._element.get_or_add_pPr()
        border = paragraph_properties.find(qn("w:pBdr"))
        if border is not None:
            paragraph_properties.remove(border)

    for style_name in ("List Bullet", "List Bullet 2"):
        style = document.styles[style_name]
        style.font.name = FONT
        style.font.size = BODY_SIZE
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.space_after = Pt(4)


def add_entry_points(document, entries):
    if not entries:
        return
    document.add_heading("Recurring entry points", level=2)
    for entry in entries:
        paragraph = document.add_paragraph()
        if entry.get("href"):
            add_hyperlink(paragraph, entry["title"], entry["href"])
        else:
            add_text(paragraph, entry["title"], bold=True)
        add_text(paragraph, f" — {entry['timing']}", bold=True)
        add_text(paragraph, f". {entry['detail']}")
        if entry.get("status"):
            add_text(paragraph, f" Status: {entry['status']}.", italic=True)
        for link in entry.get("links", []):
            add_text(paragraph, " ")
            add_hyperlink(paragraph, link["label"], link["href"])
            add_text(paragraph, ".")


def generate_document(data, process, output_path):
    operational = process["operational"]
    overview = operational["overview"]
    document = Document()
    configure_document(document)

    properties = document.core_properties
    properties.title = f"{process['acronym']}: {process['name']} — prose export"
    properties.subject = "UNYD process guide website content"
    properties.author = "UNYD process guide"

    title = document.add_paragraph(style="Title")
    add_text(title, f"{process['acronym']}: {process['name']}", bold=True)

    add_labeled_paragraph(document, "Location", process["location"])
    add_labeled_paragraph(document, "2027 date or window", process["date2027"])
    add_labeled_paragraph(document, "Date status", process["dateStatus"])
    add_labeled_paragraph(document, "Website content last verified", data["siteMeta"]["lastVerified"])
    document.add_paragraph(
        "This document contains the substantive information currently published on the UNYD process guide for this process. "
        "Official and primary links are included for checking updates. Planning estimates and recent-cycle precedents are not UN deadlines."
    )

    document.add_heading("Process overview", level=1)
    document.add_paragraph(process["summary"])
    add_labeled_paragraph(document, "Mandate", overview["mandate"])
    add_labeled_paragraph(document, "Main outputs", overview["outputs"])
    add_labeled_paragraph(document, "Working structure", overview["workingStructure"])
    add_labeled_paragraph(document, "What access depends on", overview["access"])
    add_labeled_paragraph(document, "Recurring youth relevance", overview["youthRelevance"])
    add_entry_points(document, overview.get("strategicEntryPoints", []))

    document.add_heading("UNYD participation", level=1)
    for item in process["participation"]:
        add_bullet(document, item)

    document.add_heading("Session structure", level=1)
    for item in process["structure"]:
        detail = item.get("detail", "")
        add_bullet(document, detail, label=f"{item['when']} — {item['label']}")

    document.add_heading("Preparation and negotiation calendar", level=1)
    document.add_paragraph(operational["calendarNote"])
    for milestone in (item for item in operational["milestones"] if item["status"] != "recent cycle"):
        label = f"{milestone['label']} — {milestone['start']} to {milestone['end']}"
        text = f"{milestone['detail']} Status: {milestone['status']}."
        add_bullet(document, text, label=label)

    document.add_heading("2027 file watch", level=2)
    add_labeled_paragraph(document, "Published", process["fileWatch"]["published"])
    add_labeled_paragraph(document, "Expected next", process["fileWatch"]["expected"])
    source_paragraph = document.add_paragraph()
    add_text(source_paragraph, "Official page: ", bold=True)
    add_hyperlink(source_paragraph, process["fileWatch"]["source"]["label"], process["fileWatch"]["source"]["href"])
    add_labeled_paragraph(document, "Practical route", process["fileWatch"]["route"])

    document.add_heading("Negotiations and policy work", level=1)
    add_labeled_paragraph(document, "What to follow", process["negotiations"]["focus"])
    document.add_heading("Watch for", level=2)
    for item in process["negotiations"]["watch"]:
        add_bullet(document, item)
    add_labeled_paragraph(document, "Indicative lead time", process["negotiations"]["leadTime"])
    add_labeled_paragraph(document, "Likely formal route", process["negotiations"]["route"])

    document.add_heading("Negotiation and language resources", level=2)
    for resource in data["negotiationResources"]:
        add_source(document, resource["label"], resource["href"], resource["description"])
    if process["id"] == "unga":
        resource = data["ungaResolutionWatchlist"]
        add_source(document, resource["label"], resource["href"], resource["description"])

    document.add_heading("Practical timeline", level=1)
    for item in process["timeline"]:
        paragraph = document.add_paragraph(style="List Bullet")
        add_text(paragraph, f"{item['when']} — ", bold=True)
        if item.get("href"):
            add_hyperlink(paragraph, item["action"], item["href"])
        else:
            add_text(paragraph, item["action"], bold=True)
        add_text(paragraph, f": {item['detail']}")

    verified_examples = [example for example in process["examples"] if example["state"] == "verified"]
    if verified_examples:
        document.add_heading("Examples", level=1)
        for example in verified_examples:
            paragraph = document.add_paragraph()
            add_text(paragraph, f"{example['kind']} — ", bold=True)
            if example.get("href"):
                add_hyperlink(paragraph, example["title"], example["href"])
            else:
                add_text(paragraph, example["title"], bold=True)
            add_text(paragraph, f". {example['detail']}")
    document.add_heading("Contacts", level=1)
    for contact in process["contacts"]:
        paragraph = document.add_paragraph(style="List Bullet")
        add_text(paragraph, f"{contact['label']}: ", bold=True)
        add_text(paragraph, f"{contact['role']} — ")
        add_hyperlink(paragraph, contact["email"], f"mailto:{contact['email']}")

    document.add_heading("General UN youth contacts", level=2)
    for contact in data["generalContacts"]:
        paragraph = document.add_paragraph(style="List Bullet")
        add_text(paragraph, f"{contact['label']}: ", bold=True)
        add_text(paragraph, f"{contact['role']} — ")
        add_hyperlink(paragraph, contact["email"], f"mailto:{contact['email']}")

    document.add_heading("Official and primary sources", level=1)
    for source in process["sources"]:
        add_source(document, source["label"], source["href"])

    document.save(output_path)


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: generate-process-prose-briefs.py DATA.json OUTPUT_DIRECTORY")
    data_path = Path(sys.argv[1]).resolve()
    output_directory = Path(sys.argv[2]).resolve()
    output_directory.mkdir(parents=True, exist_ok=True)
    data = json.loads(data_path.read_text(encoding="utf-8"))
    for process in data["processes"]:
        output_path = output_directory / f"{process['id']}-prose-brief.docx"
        generate_document(data, process, output_path)
        print(output_path)


if __name__ == "__main__":
    main()
